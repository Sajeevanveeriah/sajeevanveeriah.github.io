#!/usr/bin/env python3
"""Generate Sajeevan Veeriah's public two-page resume as a DOCX file."""

from __future__ import annotations

import argparse
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

INK = '111111'
MUTED = '565861'
LINE = 'CFD3DB'
PANEL = 'FFFFFF'
SOFT = 'EEF0F5'
SAPPHIRE = '001391'
VOID = '050509'
FROST = 'F8F8F8'
CONTACT = 'C7CAD2'

DEFAULT_OUTPUT = Path('public/assets/Resume_Sajeevan_Veeriah.docx')
DEFAULT_LOGO = Path('public/assets/image/20260827-Sajeevan-Veeriah-SV-Logo-Rev00.webp')


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for name, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        element = tc_mar.find(qn(f'w:{name}'))
        if element is None:
            element = OxmlElement(f'w:{name}')
            tc_mar.append(element)
        element.set(qn('w:w'), str(value))
        element.set(qn('w:type'), 'dxa')


def set_cell_borders(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge, spec in edges.items():
        tag = f'w:{edge}'
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in spec.items():
            element.set(qn(f'w:{key}'), str(value))


def remove_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(
                cell,
                top={'val': 'nil'},
                bottom={'val': 'nil'},
                start={'val': 'nil'},
                end={'val': 'nil'},
            )


def set_table_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in('w:tblLayout')
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')


def set_run(run, font: str, size: float, *, bold=False, colour=INK, italic=False) -> None:
    run.font.name = font
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn('w:ascii'), font)
    r_fonts.set(qn('w:hAnsi'), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(colour)


def style_paragraph(paragraph, *, before=0, after=0, line=1.0, align=None, keep_next=False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next
    fmt.keep_together = True
    if align is not None:
        paragraph.alignment = align


def add_text(paragraph, text: str, *, font='Liberation Sans', size=8.2, bold=False, colour=INK, italic=False):
    run = paragraph.add_run(text)
    set_run(run, font, size, bold=bold, colour=colour, italic=italic)
    return run


def add_hyperlink(paragraph, text: str, url: str, *, size=7.2, colour=CONTACT) -> None:
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), relation_id)
    run = OxmlElement('w:r')
    properties = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'), 'Liberation Sans')
    fonts.set(qn('w:hAnsi'), 'Liberation Sans')
    properties.append(fonts)
    colour_node = OxmlElement('w:color')
    colour_node.set(qn('w:val'), colour)
    properties.append(colour_node)
    size_node = OxmlElement('w:sz')
    size_node.set(qn('w:val'), str(int(size * 2)))
    properties.append(size_node)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'none')
    properties.append(underline)
    run.append(properties)
    text_node = OxmlElement('w:t')
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph, instruction: str, placeholder: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    code = OxmlElement('w:instrText')
    code.set(qn('xml:space'), 'preserve')
    code.text = instruction
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = placeholder
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, code, separate, text, end])
    set_run(run, 'Liberation Mono', 6.2, colour=MUTED)


def set_picture_alt(shape, text: str) -> None:
    shape._inline.docPr.set('name', text)
    shape._inline.docPr.set('descr', text)
    shape._inline.docPr.set('title', text)


def prepare_logo(logo_path: Path, output_path: Path) -> None:
    with Image.open(logo_path) as source:
        image = source.convert('RGBA')
        bounds = image.getchannel('A').getbbox()
        if bounds:
            image = image.crop(bounds)
        image.thumbnail((256, 256), Image.Resampling.LANCZOS)
        image = image.quantize(colors=256, method=Image.Quantize.FASTOCTREE, dither=Image.Dither.FLOYDSTEINBERG)
        image.save(output_path, 'PNG', optimize=True, compress_level=9)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(17)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.header_distance = Mm(5)
    section.footer_distance = Mm(6)

    normal = document.styles['Normal']
    normal.font.name = 'Liberation Sans'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Liberation Sans')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Liberation Sans')
    normal.font.size = Pt(8.2)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.05

    properties = document.core_properties
    properties.title = 'Sajeevan Veeriah - Public Resume'
    properties.subject = 'Robotics, Mechatronics, AI/ML & End-To-End Automation Engineer'
    properties.author = 'Sajeevan Veeriah'
    properties.last_modified_by = 'Sajeevan Veeriah'
    properties.keywords = 'robotics, mechatronics, automation, industrial IT, OT, MES, ERP, AI/ML'


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = ''
    style_paragraph(paragraph, after=0, line=1.0)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '5')
    top.set(qn('w:color'), LINE)
    p_bdr.append(top)
    p_pr.append(p_bdr)

    table = footer.add_table(rows=1, cols=2, width=Mm(180))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed(table)
    table.columns[0].width = Mm(120)
    table.columns[1].width = Mm(60)
    remove_table_borders(table)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    style_paragraph(left, after=0, line=1.0)
    style_paragraph(right, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(left, 'SAJEEVAN VEERIAH  |  PUBLIC RESUME', font='Liberation Mono', size=6.2, bold=True, colour=MUTED)
    add_text(right, 'PAGE ', font='Liberation Mono', size=6.2, bold=True, colour=MUTED)
    add_page_field(right, 'PAGE', '1')
    add_text(right, ' OF ', font='Liberation Mono', size=6.2, bold=True, colour=MUTED)
    add_page_field(right, 'NUMPAGES', '2')


def add_header(document: Document, logo_png: Path, compact=False) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed(table)
    remove_table_borders(table)
    left, right = table.rows[0].cells
    if compact:
        table.columns[0].width = Mm(147)
        table.columns[1].width = Mm(33)
        table.rows[0].height = Mm(18)
    else:
        table.columns[0].width = Mm(150)
        table.columns[1].width = Mm(30)
        table.rows[0].height = Mm(34)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for cell in (left, right):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, VOID if cell is left else PANEL)
    set_cell_margins(left, top=110 if compact else 150, start=220, bottom=110 if compact else 150, end=160)
    set_cell_margins(right, top=70, start=80, bottom=70, end=80)

    if compact:
        p = left.paragraphs[0]
        style_paragraph(p, after=0, line=1.0)
        add_text(p, 'Sajeevan Veeriah', font='Liberation Serif', size=14.5, colour=FROST)
        meta = left.add_paragraph()
        style_paragraph(meta, before=0, after=0, line=1.0)
        add_text(meta, 'ROBOTICS - MECHATRONICS - AI/ML - INDUSTRIAL IT/OT - AUTOMATION', font='Liberation Mono', size=6.0, colour=CONTACT)
        logo_width = Mm(10)
    else:
        p = left.paragraphs[0]
        style_paragraph(p, after=0, line=1.0)
        add_text(p, 'Sajeevan Veeriah', font='Liberation Serif', size=25.5, colour=FROST)
        role = left.add_paragraph()
        style_paragraph(role, after=2, line=1.0)
        add_text(role, 'Robotics, Mechatronics, AI/ML & End-To-End Automation Engineer', size=9.4, bold=True, colour=FROST)
        contact = left.add_paragraph()
        style_paragraph(contact, after=0, line=1.0)
        add_text(contact, 'Geelong, Victoria, Australia | +61 498 586 654 | ', size=7.0, colour=CONTACT)
        add_hyperlink(contact, 'sajeevanveeriah@gmail.com', 'mailto:sajeevanveeriah@gmail.com', size=7.0)
        links = left.add_paragraph()
        style_paragraph(links, after=0, line=1.0)
        add_hyperlink(links, 'sajeevanveeriah.github.io', 'https://sajeevanveeriah.github.io', size=7.0)
        add_text(links, ' | ', size=7.0, colour=CONTACT)
        add_hyperlink(links, 'linkedin.com/in/sajeevan-veeriah', 'https://www.linkedin.com/in/sajeevan-veeriah', size=7.0)
        add_text(links, ' | ', size=7.0, colour=CONTACT)
        add_hyperlink(links, 'github.com/Sajeevanveeriah', 'https://github.com/Sajeevanveeriah', size=7.0)
        logo_width = Mm(21)

    logo_paragraph = right.paragraphs[0]
    style_paragraph(logo_paragraph, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    shape = logo_paragraph.add_run().add_picture(str(logo_png), width=logo_width)
    set_picture_alt(shape, 'Sajeevan Veeriah SV portfolio monogram')


def add_section_heading(document: Document, number: str, title: str) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed(table)
    table.columns[0].width = Mm(13)
    table.columns[1].width = Mm(167)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=50, start=0, bottom=60, end=0)
        set_cell_borders(
            cell,
            bottom={'val': 'single', 'sz': '4', 'space': '4', 'color': LINE},
            top={'val': 'nil'},
            start={'val': 'nil'},
            end={'val': 'nil'},
        )
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    p = table.cell(0, 0).paragraphs[0]
    style_paragraph(p, after=0, line=1.0, keep_next=True)
    add_text(p, number, font='Liberation Mono', size=6.8, bold=True, colour=SAPPHIRE)
    p = table.cell(0, 1).paragraphs[0]
    style_paragraph(p, after=0, line=1.0, keep_next=True)
    add_text(p, title, font='Liberation Serif', size=11.7)


def add_body(document: Document, text: str, *, size=8.2, colour=INK, bold=False, after=1.2, line=1.08) -> None:
    p = document.add_paragraph()
    style_paragraph(p, after=after, line=line)
    add_text(p, text, size=size, colour=colour, bold=bold)


def add_bullet(document: Document, text: str, *, size=7.6, after=0.4) -> None:
    p = document.add_paragraph()
    style_paragraph(p, after=after, line=1.04)
    p.paragraph_format.left_indent = Mm(3.5)
    p.paragraph_format.first_line_indent = Mm(-3.5)
    add_text(p, '- ' + text, size=size)


def add_incoming(document: Document) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed(table)
    table.columns[0].width = Mm(180)
    cell = table.cell(0, 0)
    set_cell_shading(cell, SOFT)
    set_cell_margins(cell, top=100, start=220, bottom=90, end=180)
    set_cell_borders(
        cell,
        start={'val': 'single', 'sz': '18', 'color': SAPPHIRE},
        top={'val': 'single', 'sz': '4', 'color': LINE},
        bottom={'val': 'single', 'sz': '4', 'color': LINE},
        end={'val': 'single', 'sz': '4', 'color': LINE},
    )
    p = cell.paragraphs[0]
    style_paragraph(p, after=1.2, line=1.0, keep_next=True)
    add_text(p, 'INCOMING APPOINTMENT - CONTRACT PENDING', font='Liberation Mono', size=6.4, bold=True, colour=SAPPHIRE)
    p = cell.add_paragraph()
    style_paragraph(p, after=0.4, line=1.0, keep_next=True)
    add_text(p, 'Industrial IT Engineer - IT Systems / MES / ERP', font='Liberation Serif', size=11.8)
    p = cell.add_paragraph()
    style_paragraph(p, after=1.0, line=1.0, keep_next=True)
    add_text(p, 'Farm Frites Australia | Dooen, Victoria', size=7.8, bold=True)
    p = cell.add_paragraph()
    style_paragraph(p, after=1.0, line=1.02)
    add_text(p, 'Planned appointment. Employment has not commenced; the final public title and commencement date remain subject to the issued contract.', size=6.9, colour=MUTED)
    for text in (
        'Role scope includes MES implementation and production workflows, ERP integration, SCADA/PLC/historian interfaces, OT infrastructure and production-data integrity.',
        'The position also covers commissioning, validation, traceability, vendor coordination, cybersecurity, documentation and cross-functional support across Operations, Engineering, Quality, Supply Chain and IT.',
    ):
        p = cell.add_paragraph()
        style_paragraph(p, after=0.2, line=1.03)
        p.paragraph_format.left_indent = Mm(3.5)
        p.paragraph_format.first_line_indent = Mm(-3.5)
        add_text(p, '- ' + text, size=7.2)


def add_expertise(document: Document) -> None:
    items = [
        ('Industrial IT, OT and Automation', 'PLC and HMI/SCADA integration, MES exposure, production-data interfaces, field devices, drives, industrial networks, FAT, SAT and commissioning.'),
        ('Robotics and Mechatronics', 'ROS 2 Humble, Nav2, MoveIt 2, Gazebo Fortress, SLAM, EKF sensor fusion, localisation, planning, motion control and physical prototyping.'),
        ('Embedded Systems and Connectivity', 'ESP32, STM32, FreeRTOS, C/C++, PCB bring-up, CAN, UART, I2C, SPI, BLE, LoRaWAN, MQTT and Linux services.'),
        ('AI/ML, Data and Engineering Software', 'Python, scikit-learn, OpenCV, YOLO, time-series analysis, MATLAB, Simulink, TypeScript, React, Rust, SQL and CI/CD.'),
        ('Verification and Delivery', 'Requirements, instrumentation, calibrated testing, regression, fault isolation, data integrity, traceable QA, validation and technical handover.'),
        ('Manufacturing and Quality', 'GMP, food and beverage operations, traceability, batch discipline, process troubleshooting, ITP/MDR records, WHS and regulated evidence.'),
    ]
    table = document.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed(table)
    table.columns[0].width = Mm(90)
    table.columns[1].width = Mm(90)
    for index, (title, body) in enumerate(items):
        row, column = divmod(index, 2)
        cell = table.cell(row, column)
        set_cell_shading(cell, PANEL)
        set_cell_margins(cell, top=70, start=120, bottom=70, end=120)
        set_cell_borders(
            cell,
            top={'val': 'single', 'sz': '4', 'color': LINE},
            bottom={'val': 'single', 'sz': '4', 'color': LINE},
            start={'val': 'single', 'sz': '4', 'color': LINE},
            end={'val': 'single', 'sz': '4', 'color': LINE},
        )
        p = cell.paragraphs[0]
        style_paragraph(p, after=0.6, line=1.0, keep_next=True)
        add_text(p, title, size=7.2, bold=True)
        p = cell.add_paragraph()
        style_paragraph(p, after=0, line=1.01)
        add_text(p, body, size=6.65, colour=MUTED)


def add_experience(document: Document, role: str, organisation: str, period: str, context: str, bullets: list[str]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed(table)
    remove_table_borders(table)
    table.columns[0].width = Mm(134)
    table.columns[1].width = Mm(46)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
    p = table.cell(0, 0).paragraphs[0]
    style_paragraph(p, after=0, line=1.0, keep_next=True)
    add_text(p, role, font='Liberation Serif', size=9.2)
    p = table.cell(0, 1).paragraphs[0]
    style_paragraph(p, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT, keep_next=True)
    add_text(p, period, font='Liberation Mono', size=6.2, colour=MUTED)
    p = document.add_paragraph()
    style_paragraph(p, after=0, line=1.0, keep_next=True)
    add_text(p, organisation, size=7.1, bold=True)
    p = document.add_paragraph()
    style_paragraph(p, after=0.4, line=1.0, keep_next=True)
    add_text(p, context, font='Liberation Mono', size=6.25, colour=MUTED)
    for bullet_text in bullets:
        add_bullet(document, bullet_text, size=7.15, after=0.1)
    spacer = document.add_paragraph()
    style_paragraph(spacer, after=0.5, line=0.6)


def add_project(document: Document, title: str, meta: str, body: str) -> None:
    p = document.add_paragraph()
    style_paragraph(p, after=0, line=1.0, keep_next=True)
    add_text(p, title, font='Liberation Serif', size=8.9)
    p = document.add_paragraph()
    style_paragraph(p, after=0.3, line=1.0, keep_next=True)
    add_text(p, meta, font='Liberation Mono', size=6.1, colour=MUTED)
    p = document.add_paragraph()
    style_paragraph(p, after=0.8, line=1.02)
    add_text(p, body, size=7.0)


def add_labelled_line(document: Document, label: str, text: str) -> None:
    p = document.add_paragraph()
    style_paragraph(p, after=0.3, line=1.02)
    add_text(p, label, size=6.9, bold=True)
    add_text(p, text, size=6.9)


def build(output: Path, logo: Path) -> None:
    if not logo.is_file():
        raise FileNotFoundError(f'Portfolio monogram not found: {logo}')
    output.parent.mkdir(parents=True, exist_ok=True)
    logo_png = output.with_name('.resume-monogram.png')
    prepare_logo(logo, logo_png)

    document = Document()
    configure_document(document)
    add_footer(document.sections[0])

    add_header(document, logo_png)
    add_section_heading(document, '01', 'Professional Profile')
    add_body(
        document,
        'I engineer complete systems across the physical, control, digital and operational boundary, then verify them as one working system. My practice spans robotics, industrial automation, embedded systems, industrial IT/OT, AI/ML and engineering software from requirements through integration, commissioning and handover.',
        size=8.0,
        after=0.5,
    )

    add_section_heading(document, '02', 'Incoming Appointment')
    add_incoming(document)

    add_section_heading(document, '03', 'Core Expertise')
    add_expertise(document)

    add_section_heading(document, '04', 'Professional Experience')
    add_experience(
        document,
        'Automation and Controls Engineer',
        'Process automation systems integrator',
        'Jan 2026 - Jun 2026',
        'Regulated pharmaceutical, biotechnology and food production',
        [
            'Delivered PLC logic, HMI/SCADA, field-device, drive, MES and production-data integration under GMP controls.',
            'Migrated validated application content from iFIX to PVI+, completed functional checks, and supported FAT, SAT, commissioning, handover and MiR mobile-robot fault tracing.',
        ],
    )
    add_experience(
        document,
        'Product Development Test Engineer (Contract)',
        'Global automotive OEM via engineering consultancy',
        'Oct 2025 - Jan 2026',
        'Vehicle software integration and ADAS product development',
        [
            'Validated vehicle-software integration and ADAS behaviour through feature-vehicle, breadboard and regression testing.',
            'Instrumented test vehicles and used Vector CANoe and CANalyzer to capture CAN and CAN FD evidence for fault isolation and software-readiness verification.',
        ],
    )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_header(document, logo_png, compact=True)
    add_section_heading(document, '04', 'Professional Experience Continued')
    add_experience(
        document,
        'Technical Officer and Quality Assurance',
        'ABMARC',
        'Jul 2024 - Aug 2025',
        'Automotive testing, energy, emissions and compliance engineering',
        [
            'Delivered ADR and EURO emissions work, EV/PHEV range testing and vehicle-systems evaluation using calibrated instrumentation, data acquisition and CAN tools.',
            'Maintained traceable QA, safety, regulatory and test records for certification review and audit.',
        ],
    )
    add_experience(
        document,
        'Consultant Engineer, IoT and Projects Administrator',
        'DuxTel',
        'Feb 2024 - Aug 2024',
        'Networking, low-power telemetry and remote asset visibility',
        [
            'Built field-to-dashboard solutions with ESP32, sensors, LoRaWAN AU915, MQTT, ChirpStack, InfluxDB and Grafana.',
            'Supported CAN and GPS capture, custom PCB work, MikroTik connectivity and Linux services from field trial through handover.',
        ],
    )
    add_experience(
        document,
        'Manufacturing, Commissioning and Quality Foundation',
        'IDL Australia | Carbon Revolution | Thornton Engineering Australia',
        '2018 - 2024',
        'Food and beverage, carbon-fibre and structural-steel production',
        [
            'Built practical experience in equipment operation, changeovers, safe production, first-response fault recovery, process discipline, batch quality and traceability.',
            'Supported drawing review, ITP/MDR documentation, material traceability, QA inspection, equipment trials and commissioning checks across workshop and production environments.',
        ],
    )

    add_section_heading(document, '05', 'Selected Engineering Projects')
    add_project(
        document,
        'Open Industrial Automation',
        '2026 | Public industrial reference and simulation platform',
        'Built a deterministic mixing, dosing and CIP reference process with an operator HMI, Engineering Studio, typed project model, traceability, cybersecurity boundaries and automated verification gates.',
    )
    add_project(
        document,
        'Autonomous Navigation Rover on ROS 2',
        '2024 - 2025 | Hardware build with simulation-validated autonomy',
        'Built the differential-drive platform and integrated LiDAR, IMU, SLAM, EKF state estimation, Nav2 planning, control and recovery across independently testable ROS 2 nodes.',
    )
    add_project(
        document,
        'ESP32 Clinical Ataxia Assessment Device',
        '2025 | Honours capstone and assessed embedded prototype',
        'Designed the device, custom PCB and enclosure, implemented deterministic 100 Hz acquisition, Bluetooth live display and CSV/PDF reporting, then checked accuracy, reversal, drift and temperature behaviour against reference instruments.',
    )

    add_section_heading(document, '06', 'Education')
    add_body(document, 'Bachelor of Mechatronics Engineering (Honours), Distinction', size=7.25, bold=True, after=0.1)
    add_body(document, 'Deakin University, Geelong | 2025 | Honours capstone: ESP32 clinical ataxia assessment device.', size=6.75, after=0.5)
    add_body(document, 'Higher National Diploma in Mechatronics, Robotics and Automation Engineering, Distinction', size=7.25, bold=True, after=0.1)
    add_body(document, 'Cardiff Metropolitan University, United Kingdom | 2016.', size=6.75, after=0.5)

    add_section_heading(document, '07', 'Professional Membership')
    add_body(document, 'Member, Engineers Australia', size=7.0, bold=True, after=0.1)

    add_section_heading(document, '08', 'Professional Development')
    add_body(document, 'Lean Six Sigma Foundation | JIRA and Agile | KAIZEN | Industrial Automation and IIoT | AI/ML | CAD', size=6.8, after=0.1)

    add_section_heading(document, '09', 'Languages and Additional Information')
    add_labelled_line(document, 'Languages: ', 'English, Tamil and Sinhala.')
    add_labelled_line(document, 'Additional: ', 'Current Victorian driver licence | Community leadership through Newcomb and District Cricket Club | Engineering mentoring and peer support.')

    for table in document.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cannot_split = OxmlElement('w:cantSplit')
            tr_pr.append(cannot_split)

    document.save(output)
    logo_png.unlink(missing_ok=True)

    from zipfile import ZipFile
    with ZipFile(output) as archive:
        names = set(archive.namelist())
        for required in ('[Content_Types].xml', 'word/document.xml', 'word/media/image1.png'):
            if required not in names:
                raise RuntimeError(f'Missing required DOCX part: {required}')
        xml = archive.read('word/document.xml').decode('utf-8')
    for phrase in ('Sajeevan Veeriah', 'Farm Frites Australia', 'CONTRACT PENDING', 'Member, Engineers Australia'):
        if phrase not in xml:
            raise RuntimeError(f'Missing required resume phrase: {phrase}')
    for forbidden in ('Ford Motor Company', 'Invenio', 'JAG Process Solutions', '\u2013', '\u2014'):
        if forbidden in xml:
            raise RuntimeError(f'Forbidden public resume content: {forbidden}')


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument('--output', type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument('--logo', type=Path, default=DEFAULT_LOGO)
    args = parser.parse_args()
    build(args.output, args.logo)
    print(args.output)


if __name__ == '__main__':
    main()
